import sys
if sys.version_info[0] != 3 or sys.version_info[1] < 10:
    raise SystemExit('Python version too low, at least Python 3.10 can use this module')

from tkinter import *
from docx import *
from docx.shared import *

def add_text(text, attribute=Document(), mode=('normal', 'attribute', None), paragraph=(True, 'normal', None)):  # mode => text - return - argument ; paragraph => new - mode - argument
    if mode[1] == 'attribute':
        if paragraph[0] == True:
            match paragraph[1]:
                case 'normal':
                    p = attribute.add_paragraph('')
                case 'body':
                    p = attribute.add_paragraph('', style='Body Text')
                case 'bullet':
                    p = attribute.add_paragraph('', style='List Bullet')
                case 'number':
                    p = attribute.add_paragraph('', style='List Number')
                case _:
                    raise AttributeError('\'add_text - paragraph - mode\' object has no attribute \'%s\'' % paragraph[1])
        else:
            p = argument
        # Paragraph
        match mode[0]:
            case 'normal':
                p.add_run(text)
            case'normal':
                p.add_run(text)
            case 'bold':
                p.add_run(text).bold = True
            case 'italic':
                p.add_run(text).italic = True
            case 'underline':
                p.add_run(text).underline = True
            case 'deleteline':
                p.add_run(text).deleteline = True
            case 'outline':
                p.add_run(text).outline = True
            case 'imprint':
                p.add_run(text).imprint = True
            case 'emboss':
                p.add_run(text).emboss = True
            case 'double_strike':
                p.add_run(text).double_strike = True
            case 'rtl':
                p.add_run(text).rtl = True
            case 'shadow':
                p.add_run(text).shadow = True
            case 'size':
                p.add_run(text).size = Pt(mode[2])
            case 'subscript':
                p.add_run(text).subscript = True
            case 'superscript':
                p.add_run(text).subscript = True
            case _:
                raise AttributeError('\'add_text - run - mode\'object has no attribute \'%s\'' % mode[0])
        # Add Run
        return attribute
    elif mode[1] == 'list':
        p = []
        match mode[0]:
            case 'normal':
                p.append(text)
            case 'bold':
                p.append(['bold', text])
            case 'italic':
                p.append(['italic', text])
            case 'underline':
                p.append(['underline', text])
            case 'deleteline':
                p.append(['deleteline', text])
            case 'outline':
                p.append(['outline', text])
            case 'imprint':
                p.append(['imprint', text])
            case 'emboss':
                p.append(['emboss', text])
            case 'double_strike':
                p.append(['double strike', text])
            case 'rtl':
                p.append(['rtl', text])
            case 'shadow':
                p.append(['shadow', text])
            case 'size':
                p.append(['size', text, Pt(mode[2])])
            case 'subscript':
                p.append(['subscript', text])
            case 'superscript':
                p.append(['superscript', text])
            case _:
                raise AttributeError('\'add_text - run - mode\'object has no attribute \'%s\'' % mode[0])
        return p
    else:
        raise AttributeError('\'add_text\' has no attribute \'%s\'' % mode[1])

def add_heading(text, attribute=Document(), mode=(None), heading=(True, 'normal', None)):  # mode => argumet ; heading => new - mode - argument
    if heading[0] == True:
        match heading[1]:
            case 'normal':
                h = attribute.add_heading(text, 0)
            case _:
                raise AttributeError('\'add_heading - heading - mode\' object has no attribute \'%s\'' % heading[1])
    return attribute
