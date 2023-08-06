# -*- coding:utf-8  -*-

from tkinter import *
from docx import *
from docx.shared import *

def add_text(text, attribute=Document(), mode=('normal', 'attribute', None), paragraph=(True, 'normal', None)) -> 'mode = ( text, return, argument ) ; paragraph = ( new, mode, argument )':
    if mode[1] == 'attribute':
        if paragraph[0] == True:
            if paragraph[1] == 'normal':
                p = attribute.add_paragraph('')
            elif paragraph[1] == 'body':
                p = attribute.add_paragraph('', style='Body Text')
            elif paragraph[1] == 'bullet':
                p = attribute.add_paragraph('', style='List Bullet')
            elif paragraph[1] == 'number':
                p = attribute.add_paragraph('', style='List Number')
            else:
                raise AttributeError('\'add_text - paragraph - mode\' object has no attribute \'%s\'' % paragraph[1])
        else:
            p = argument
        # Paragraph
        if mode[0] == 'normal':
            p.add_run(text)
        elif mode[0] =='normal':
            p.add_run(text)
        elif mode[0] == 'bold':
            p.add_run(text).bold = True
        elif mode[0] == 'italic':
            p.add_run(text).italic = True
        elif mode[0] == 'underline':
            p.add_run(text).underline = True
        elif mode[0] == 'deleteline':
            p.add_run(text).deleteline = True
        elif mode[0] == 'outline':
            p.add_run(text).outline = True
        elif mode[0] == 'imprint':
            p.add_run(text).imprint = True
        elif mode[0] == 'emboss':
            p.add_run(text).emboss = True
        elif mode[0] == 'double_strike':
            p.add_run(text).double_strike = True
        elif mode[0] == 'rtl':
            p.add_run(text).rtl = True
        elif mode[0] == 'shadow':
            p.add_run(text).shadow = True
        elif mode[0] == 'size':
            p.add_run(text).size = Pt(mode[2])
        elif mode[0] == 'subscript':
            p.add_run(text).subscript = True
        elif mode[0] == 'superscript':
            p.add_run(text).subscript = True
        else:
            raise AttributeError('\'add_text - run - mode\'object has no attribute \'%s\'' % mode[0])
        # Add Run
        return attribute
    elif mode[1] == 'list':
        p = []
        if mode[0] == 'normal':
            p.append(text)
        elif mode[0] == 'bold':
            p.append(['bold', text])
        elif mode[0] == 'italic':
            p.append(['italic', text])
        elif mode[0] == 'underline':
            p.append(['underline', text])
        elif mode[0] == 'deleteline':
            p.append(['deleteline', text])
        elif mode[0] == 'outline':
            p.append(['outline', text])
        elif mode[0] == 'imprint':
            p.append(['imprint', text])
        elif mode[0] == 'emboss':
            p.append(['emboss', text])
        elif mode[0] == 'double_strike':
            p.append(['double strike', text])
        elif mode[0] == 'rtl':
            p.append(['rtl', text])
        elif mode[0] == 'shadow':
            p.append(['shadow', text])
        elif mode[0] == 'size':
            p.append(['size', text, Pt(mode[2])])
        elif mode[0] == 'subscript':
            p.append(['subscript', text])
        elif mode[0] == 'superscript':
            p.append(['superscript', text])
        else:
            raise AttributeError('\'add_text - run - mode\'object has no attribute \'%s\'' % mode[0])
        return p
    else:
        raise AttributeError('\'add_text\' has no attribute \'%s\'' % mode[1])

def add_heading(text, attribute=Document(), mode=(None), heading=(True, 'normal', None)) -> 'mode = ( argumet ) ; heading = ( new, mode, argument )':
    if heading[0] == True:
        if heading[1] == 'normal':
            h = attribute.add_heading(text, 0)
        else:
            raise AttributeError('\'add_heading - heading - mode\' object has no attribute \'%s\'' % heading[1])
    return attribute

